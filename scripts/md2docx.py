#!/usr/bin/env python3
"""Ensayo_Final.md -> .docx con formato académico (TNR 12, justificado,
griego + negritas/cursivas). Ejecutar desde la raíz del repo."""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

SRC = "ensayo/Ensayo_Final.md"
OUT = "ensayo/Ensayo_Final.docx"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for section in doc.sections:
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

def add_runs(paragraph, text, base_bold=False, size=None):
    pos = 0
    def sty(r):
        if base_bold: r.bold = True
        if size: r.font.size = Pt(size)
    for m in INLINE.finditer(text):
        if m.start() > pos: sty(paragraph.add_run(text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
            if size: r.font.size = Pt(size)
        else:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
            if base_bold: r.bold = True
            if size: r.font.size = Pt(size)
        pos = m.end()
    if pos < len(text): sty(paragraph.add_run(text[pos:]))

with open(SRC, encoding="utf-8") as fh:
    lines = fh.read().split("\n")

for raw in lines:
    s = raw.strip()
    if s == "" or s == "***": continue
    if s.startswith("# "):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14); add_runs(p, s[2:], True, 14)
    elif s.startswith("### "):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
        add_runs(p, s[4:], True, 12)
    elif s.startswith("## "):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
        add_runs(p, s[3:], True, 13)
    elif s.startswith("* "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_after = Pt(4); add_runs(p, s[2:])
    else:
        add_runs(doc.add_paragraph(), s)

doc.save(OUT)
print(f"OK -> {OUT}")
